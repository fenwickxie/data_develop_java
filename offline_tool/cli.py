import json
from pathlib import Path
from typing import Optional

import requests
import typer
import yaml

from core.data_processing.candata import CanData
from core.data_processing.candecode import load_config_from_yaml

app = typer.Typer(help="Offline helper for download/compute/upload flows.")


def load_config(config_path: Path) -> dict:
    if config_path.exists():
        with config_path.open("r", encoding="utf-8") as fh:
            return yaml.safe_load(fh) or {}
    return {}


@app.command()
def download(
    file_id: str = typer.Argument(..., help="ID of the raw file"),
    signed_url: str = typer.Argument(..., help="Pre-signed URL from backend"),
    output_dir: Path = typer.Option(Path("downloads"), help="Where to save the file")
):
    output_dir.mkdir(parents=True, exist_ok=True)
    target = output_dir / f"{file_id}"
    typer.echo(f"Downloading {file_id} -> {target}")
    with requests.get(signed_url, stream=True, timeout=30) as resp:
        resp.raise_for_status()
        with target.open("wb") as fh:
            for chunk in resp.iter_content(chunk_size=1024 * 1024):
                if chunk:
                    fh.write(chunk)
    typer.echo("Download complete")


@app.command()
def compute(
    input_path: Path = typer.Argument(..., exists=True, readable=True, help="Local file to analyze (.csv aggregated CAN or BLF/ASC raw)"),
    output_path: Path = typer.Option(Path("metrics/metrics.json"), help="Where to write computed metrics"),
    dbc: Optional[Path] = typer.Option(None, help="DBC file for BLF/ASC decode"),
    step: float = typer.Option(0.02, help="Raster step when decoding BLF/ASC")
):
    output_path.parent.mkdir(parents=True, exist_ok=True)

    suffix = input_path.suffix.lower()
    if suffix == ".csv":
        typer.echo("Running CanData metrics extraction for CSV...")
        can_data = CanData(str(input_path))
        metrics = can_data.get_all_metrics()
        output_path.write_text(json.dumps(metrics.all_metrics, indent=2, default=str), encoding="utf-8")
        typer.echo(f"Wrote metrics -> {output_path}")
    elif suffix in {".blf", ".asc"} and dbc:
        typer.echo("Decoding BLF/ASC via core.candecode with full raster processing...")
        from core.data_processing.candecode import process_candecode_from_config
        
        cfg = {
            "dbc_path": str(dbc),
            "can_data_path": str(input_path),
            "output_dir": str(output_path.parent / "decoded"),
            "step": step,
            "save_formats": [".parquet"],
            "time_from_zero": False,
        }
        config_yaml = create_tmp_cfg(cfg)
        
        try:
            decoded_data = process_candecode_from_config(config_yaml)
            # Generate basic metrics from decoded data
            metrics_output = {
                "note": "BLF/ASC decode complete",
                "signals_decoded": len(decoded_data),
                "config": cfg
            }
            output_path.write_text(json.dumps(metrics_output, indent=2), encoding="utf-8")
            typer.echo(f"BLF/ASC decode complete. Decoded data in {cfg['output_dir']}, metrics -> {output_path}")
        except Exception as e:
            typer.echo(f"Decode error: {e}")
            metrics_payload = {
                "error": str(e),
                "note": "BLF/ASC decode failed"
            }
            output_path.write_text(json.dumps(metrics_payload, indent=2), encoding="utf-8")
    else:
        metrics_payload = {
            "metric_name": "placeholder",
            "metric_value": {"note": "BLF/ASC parsing to be wired"},
            "metric_type": "demo"
        }
        output_path.write_text(json.dumps(metrics_payload, indent=2), encoding="utf-8")
        typer.echo(f"Wrote placeholder metrics -> {output_path}")


def create_tmp_cfg(cfg: dict) -> Path:
    tmp = Path(".candecode.tmp.yaml")
    tmp.write_text(yaml.safe_dump(cfg, allow_unicode=True), encoding="utf-8")
    return tmp


@app.command()
def upload(
    dataset_id: str = typer.Argument(..., help="Dataset identifier"),
    file_id: str = typer.Argument(..., help="File identifier"),
    metrics_file: Path = typer.Argument(..., exists=True, readable=True, help="Metrics JSON to upload"),
    base_url: Optional[str] = typer.Option(None, envvar="API_BASE_URL", help="Backend base URL"),
    auth_token: Optional[str] = typer.Option(None, envvar="API_AUTH_TOKEN", help="Bearer token"),
    config_path: Path = typer.Option(Path("config.yaml"), help="Config file with base_url/auth_token")
):
    config = load_config(config_path)
    api_base = base_url or config.get("base_url") or "http://localhost:8080"
    token = auth_token or config.get("auth_token")

    payload = json.loads(metrics_file.read_text(encoding="utf-8"))
    payload.update({"datasetId": dataset_id, "fileId": file_id})

    headers = {"Content-Type": "application/json"}
    if token:
        headers["Authorization"] = f"Bearer {token}"

    url = f"{api_base}/api/metrics"
    typer.echo(f"POST {url}")
    resp = requests.post(url, headers=headers, json=payload, timeout=30)
    if resp.status_code >= 400:
        typer.echo(f"Upload failed: {resp.status_code} {resp.text}")
        raise typer.Exit(code=1)

    typer.echo("Upload complete")


@app.command()
def generate_chart(
    data_path: Path = typer.Argument(..., exists=True, readable=True, help="Decoded parquet or CSV data"),
    output_dir: Path = typer.Option(Path("charts"), help="Where to save chart images"),
    signal_columns: Optional[str] = typer.Option(None, help="Comma-separated signal columns to plot")
):
    """Generate time-series charts from decoded CAN data."""
    import pandas as pd
    import matplotlib.pyplot as plt
    
    output_dir.mkdir(parents=True, exist_ok=True)
    typer.echo(f"Loading data from {data_path}...")
    
    if data_path.suffix == ".parquet":
        df = pd.read_parquet(data_path)
    elif data_path.suffix == ".csv":
        df = pd.read_csv(data_path)
    else:
        typer.echo("Unsupported file type. Use .parquet or .csv")
        raise typer.Exit(code=1)
    
    columns = signal_columns.split(",") if signal_columns else df.columns[1:6].tolist()
    
    for col in columns:
        if col not in df.columns:
            typer.echo(f"Warning: Column {col} not found, skipping...")
            continue
        
        plt.figure(figsize=(12, 6))
        plt.plot(df.index if 'timestamps' not in df.columns else df['timestamps'], df[col])
        plt.title(f"Signal: {col}")
        plt.xlabel("Time")
        plt.ylabel(col)
        plt.grid(True, alpha=0.3)
        
        chart_path = output_dir / f"{col}.png"
        plt.savefig(chart_path, dpi=150, bbox_inches='tight')
        plt.close()
        typer.echo(f"  Generated {chart_path}")
    
    typer.echo(f"Charts saved to {output_dir}")


@app.command()
def generate_report(
    metrics_path: Path = typer.Argument(..., exists=True, readable=True, help="Metrics JSON file"),
    charts_dir: Path = typer.Option(Path("charts"), help="Directory with chart images"),
    output: Path = typer.Option(Path("report/analysis_report.docx"), help="Output Word document path")
):
    """Generate analysis report with metrics and charts."""
    from docx import Document
    from docx.shared import Inches
    
    output.parent.mkdir(parents=True, exist_ok=True)
    
    doc = Document()
    doc.add_heading("CAN Data Analysis Report", 0)
    
    # Add metrics section
    doc.add_heading("Metrics Summary", 1)
    metrics = json.loads(metrics_path.read_text(encoding="utf-8"))
    for key, value in metrics.items():
        doc.add_paragraph(f"{key}: {value}")
    
    # Add charts section
    doc.add_heading("Signal Visualizations", 1)
    if charts_dir.exists():
        for chart in sorted(charts_dir.glob("*.png")):
            doc.add_heading(chart.stem, 2)
            doc.add_picture(str(chart), width=Inches(6))
    
    doc.save(str(output))
    typer.echo(f"Report generated: {output}")


if __name__ == "__main__":
    app()
