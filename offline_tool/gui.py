#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
CAN 数据分析工具 - 桌面 GUI 应用
使用 PySide6 实现图形界面，集成 CLI 所有功能
"""

import sys
import json
from pathlib import Path
from typing import Optional

from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QTabWidget, QPushButton, QLabel, QLineEdit, QTextEdit,
    QFileDialog, QProgressBar, QComboBox, QSpinBox, QDoubleSpinBox,
    QCheckBox, QGroupBox, QFormLayout, QMessageBox
)
from PySide6.QtCore import QThread, Signal, Qt
from PySide6.QtGui import QFont, QIcon

import requests
import yaml

from core.data_processing.candata import CanData
from core.data_processing.candecode import process_candecode_from_config


class WorkerThread(QThread):
    """后台工作线程，避免阻塞 UI"""
    progress = Signal(int)
    log = Signal(str)
    finished = Signal(bool, str)
    
    def __init__(self, task_func, *args, **kwargs):
        super().__init__()
        self.task_func = task_func
        self.args = args
        self.kwargs = kwargs
    
    def run(self):
        try:
            self.task_func(*self.args, **self.kwargs, 
                          progress_callback=self.progress.emit,
                          log_callback=self.log.emit)
            self.finished.emit(True, "任务完成")
        except Exception as e:
            self.finished.emit(False, f"错误: {str(e)}")


class DownloadTab(QWidget):
    """下载标签页"""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.init_ui()
    
    def init_ui(self):
        layout = QVBoxLayout()
        
        # 输入区域
        form_layout = QFormLayout()
        
        self.file_id_input = QLineEdit()
        self.file_id_input.setPlaceholderText("例如: file-001")
        form_layout.addRow("文件 ID:", self.file_id_input)
        
        self.url_input = QLineEdit()
        self.url_input.setPlaceholderText("粘贴后端返回的签名 URL")
        form_layout.addRow("签名 URL:", self.url_input)
        
        output_layout = QHBoxLayout()
        self.output_dir_input = QLineEdit("downloads")
        output_btn = QPushButton("浏览...")
        output_btn.clicked.connect(self.browse_output_dir)
        output_layout.addWidget(self.output_dir_input)
        output_layout.addWidget(output_btn)
        form_layout.addRow("保存目录:", output_layout)
        
        layout.addLayout(form_layout)
        
        # 操作按钮
        self.download_btn = QPushButton("开始下载")
        self.download_btn.clicked.connect(self.start_download)
        layout.addWidget(self.download_btn)
        
        # 进度条
        self.progress_bar = QProgressBar()
        layout.addWidget(self.progress_bar)
        
        # 日志输出
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setMaximumHeight(150)
        layout.addWidget(QLabel("日志:"))
        layout.addWidget(self.log_text)
        
        layout.addStretch()
        self.setLayout(layout)
    
    def browse_output_dir(self):
        directory = QFileDialog.getExistingDirectory(self, "选择保存目录")
        if directory:
            self.output_dir_input.setText(directory)
    
    def start_download(self):
        file_id = self.file_id_input.text().strip()
        signed_url = self.url_input.text().strip()
        output_dir = Path(self.output_dir_input.text())
        
        if not file_id or not signed_url:
            QMessageBox.warning(self, "输入错误", "请填写文件 ID 和签名 URL")
            return
        
        self.download_btn.setEnabled(False)
        self.log_text.clear()
        self.log_text.append(f"开始下载: {file_id}")
        
        output_dir.mkdir(parents=True, exist_ok=True)
        target = output_dir / file_id
        
        try:
            self.log_text.append(f"目标路径: {target}")
            with requests.get(signed_url, stream=True, timeout=30) as resp:
                resp.raise_for_status()
                total_size = int(resp.headers.get('content-length', 0))
                downloaded = 0
                
                with target.open("wb") as fh:
                    for chunk in resp.iter_content(chunk_size=1024 * 1024):
                        if chunk:
                            fh.write(chunk)
                            downloaded += len(chunk)
                            if total_size > 0:
                                progress = int(downloaded * 100 / total_size)
                                self.progress_bar.setValue(progress)
            
            self.log_text.append("✓ 下载完成")
            QMessageBox.information(self, "成功", f"文件已保存到: {target}")
        except Exception as e:
            self.log_text.append(f"✗ 下载失败: {e}")
            QMessageBox.critical(self, "错误", str(e))
        finally:
            self.download_btn.setEnabled(True)
            self.progress_bar.setValue(0)


class ComputeTab(QWidget):
    """计算标签页"""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.init_ui()
    
    def init_ui(self):
        layout = QVBoxLayout()
        
        # 输入文件
        input_group = QGroupBox("输入文件")
        input_layout = QVBoxLayout()
        
        file_layout = QHBoxLayout()
        self.input_file = QLineEdit()
        self.input_file.setPlaceholderText("选择 CSV/BLF/ASC 文件")
        browse_btn = QPushButton("浏览...")
        browse_btn.clicked.connect(self.browse_input_file)
        file_layout.addWidget(self.input_file)
        file_layout.addWidget(browse_btn)
        input_layout.addLayout(file_layout)
        
        input_group.setLayout(input_layout)
        layout.addWidget(input_group)
        
        # 配置选项
        config_group = QGroupBox("配置选项")
        config_layout = QFormLayout()
        
        output_layout = QHBoxLayout()
        self.output_file = QLineEdit("metrics/metrics.json")
        output_browse_btn = QPushButton("浏览...")
        output_browse_btn.clicked.connect(self.browse_output_file)
        output_layout.addWidget(self.output_file)
        output_layout.addWidget(output_browse_btn)
        config_layout.addRow("输出文件:", output_layout)
        
        dbc_layout = QHBoxLayout()
        self.dbc_file = QLineEdit()
        self.dbc_file.setPlaceholderText("BLF/ASC 需要 DBC 文件")
        dbc_browse_btn = QPushButton("浏览...")
        dbc_browse_btn.clicked.connect(self.browse_dbc_file)
        dbc_layout.addWidget(self.dbc_file)
        dbc_layout.addWidget(dbc_browse_btn)
        config_layout.addRow("DBC 文件:", dbc_layout)
        
        self.step_input = QDoubleSpinBox()
        self.step_input.setRange(0.001, 1.0)
        self.step_input.setValue(0.02)
        self.step_input.setDecimals(3)
        self.step_input.setSingleStep(0.01)
        config_layout.addRow("采样步长:", self.step_input)
        
        config_group.setLayout(config_layout)
        layout.addWidget(config_group)
        
        # 操作按钮
        self.compute_btn = QPushButton("开始计算")
        self.compute_btn.clicked.connect(self.start_compute)
        layout.addWidget(self.compute_btn)
        
        # 进度条
        self.progress_bar = QProgressBar()
        layout.addWidget(self.progress_bar)
        
        # 日志输出
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setMaximumHeight(150)
        layout.addWidget(QLabel("日志:"))
        layout.addWidget(self.log_text)
        
        layout.addStretch()
        self.setLayout(layout)
    
    def browse_input_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "选择输入文件", "", 
            "CAN 数据文件 (*.csv *.blf *.asc);;所有文件 (*.*)"
        )
        if file_path:
            self.input_file.setText(file_path)
    
    def browse_output_file(self):
        file_path, _ = QFileDialog.getSaveFileName(
            self, "选择输出文件", "", "JSON 文件 (*.json)"
        )
        if file_path:
            self.output_file.setText(file_path)
    
    def browse_dbc_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "选择 DBC 文件", "", "DBC 文件 (*.dbc);;所有文件 (*.*)"
        )
        if file_path:
            self.dbc_file.setText(file_path)
    
    def start_compute(self):
        input_path = Path(self.input_file.text())
        output_path = Path(self.output_file.text())
        dbc_path = Path(self.dbc_file.text()) if self.dbc_file.text() else None
        step = self.step_input.value()
        
        if not input_path.exists():
            QMessageBox.warning(self, "输入错误", "请选择有效的输入文件")
            return
        
        self.compute_btn.setEnabled(False)
        self.log_text.clear()
        self.progress_bar.setValue(0)
        
        output_path.parent.mkdir(parents=True, exist_ok=True)
        suffix = input_path.suffix.lower()
        
        try:
            if suffix == ".csv":
                self.log_text.append("正在处理 CSV 文件...")
                self.progress_bar.setValue(30)
                can_data = CanData(str(input_path))
                metrics = can_data.get_all_metrics()
                self.progress_bar.setValue(80)
                output_path.write_text(
                    json.dumps(metrics.all_metrics, indent=2, default=str), 
                    encoding="utf-8"
                )
                self.progress_bar.setValue(100)
                self.log_text.append(f"✓ 指标已保存到: {output_path}")
                
            elif suffix in {".blf", ".asc"} and dbc_path:
                self.log_text.append("正在解码 BLF/ASC 文件...")
                self.progress_bar.setValue(10)
                
                cfg = {
                    "dbc_path": str(dbc_path),
                    "can_data_path": str(input_path),
                    "output_dir": str(output_path.parent / "decoded"),
                    "step": step,
                    "save_formats": [".parquet"],
                    "time_from_zero": False,
                }
                
                tmp_cfg = Path(".candecode.tmp.yaml")
                tmp_cfg.write_text(yaml.safe_dump(cfg, allow_unicode=True), encoding="utf-8")
                
                self.progress_bar.setValue(30)
                decoded_count = process_candecode_from_config(tmp_cfg)
                self.progress_bar.setValue(80)
                
                metrics_output = {
                    "note": "BLF/ASC decode complete",
                    "signals_decoded": decoded_count,
                    "config": cfg
                }
                output_path.write_text(json.dumps(metrics_output, indent=2), encoding="utf-8")
                self.progress_bar.setValue(100)
                self.log_text.append(f"✓ 解码完成，信号数: {decoded_count}")
                self.log_text.append(f"✓ 指标已保存到: {output_path}")
                
            else:
                raise ValueError("不支持的文件类型或缺少 DBC 文件")
            
            QMessageBox.information(self, "成功", "计算完成")
            
        except Exception as e:
            self.log_text.append(f"✗ 错误: {e}")
            QMessageBox.critical(self, "错误", str(e))
        finally:
            self.compute_btn.setEnabled(True)


class UploadTab(QWidget):
    """上传标签页"""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.init_ui()
    
    def init_ui(self):
        layout = QVBoxLayout()
        
        # 配置区域
        config_group = QGroupBox("API 配置")
        config_layout = QFormLayout()
        
        self.base_url_input = QLineEdit("http://localhost:8080")
        config_layout.addRow("API 地址:", self.base_url_input)
        
        self.token_input = QLineEdit()
        self.token_input.setPlaceholderText("Bearer token（可选）")
        config_layout.addRow("认证 Token:", self.token_input)
        
        config_group.setLayout(config_layout)
        layout.addWidget(config_group)
        
        # 数据区域
        data_group = QGroupBox("上传数据")
        data_layout = QFormLayout()
        
        self.dataset_id_input = QLineEdit()
        self.dataset_id_input.setPlaceholderText("例如: dataset-001")
        data_layout.addRow("Dataset ID:", self.dataset_id_input)
        
        self.file_id_input = QLineEdit()
        self.file_id_input.setPlaceholderText("例如: file-001")
        data_layout.addRow("File ID:", self.file_id_input)
        
        metrics_layout = QHBoxLayout()
        self.metrics_file_input = QLineEdit()
        metrics_browse_btn = QPushButton("浏览...")
        metrics_browse_btn.clicked.connect(self.browse_metrics_file)
        metrics_layout.addWidget(self.metrics_file_input)
        metrics_layout.addWidget(metrics_browse_btn)
        data_layout.addRow("指标文件:", metrics_layout)
        
        data_group.setLayout(data_layout)
        layout.addWidget(data_group)
        
        # 操作按钮
        self.upload_btn = QPushButton("上传")
        self.upload_btn.clicked.connect(self.start_upload)
        layout.addWidget(self.upload_btn)
        
        # 进度条
        self.progress_bar = QProgressBar()
        layout.addWidget(self.progress_bar)
        
        # 日志输出
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setMaximumHeight(150)
        layout.addWidget(QLabel("日志:"))
        layout.addWidget(self.log_text)
        
        layout.addStretch()
        self.setLayout(layout)
    
    def browse_metrics_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "选择指标文件", "", "JSON 文件 (*.json);;所有文件 (*.*)"
        )
        if file_path:
            self.metrics_file_input.setText(file_path)
    
    def start_upload(self):
        base_url = self.base_url_input.text().strip()
        token = self.token_input.text().strip()
        dataset_id = self.dataset_id_input.text().strip()
        file_id = self.file_id_input.text().strip()
        metrics_file = Path(self.metrics_file_input.text())
        
        if not all([base_url, dataset_id, file_id]) or not metrics_file.exists():
            QMessageBox.warning(self, "输入错误", "请填写所有必填字段")
            return
        
        self.upload_btn.setEnabled(False)
        self.log_text.clear()
        self.progress_bar.setValue(0)
        
        try:
            self.log_text.append("正在读取指标文件...")
            payload = json.loads(metrics_file.read_text(encoding="utf-8"))
            payload.update({"datasetId": dataset_id, "fileId": file_id})
            self.progress_bar.setValue(30)
            
            headers = {"Content-Type": "application/json"}
            if token:
                headers["Authorization"] = f"Bearer {token}"
            
            url = f"{base_url}/api/metrics"
            self.log_text.append(f"POST {url}")
            self.progress_bar.setValue(50)
            
            resp = requests.post(url, headers=headers, json=payload, timeout=30)
            self.progress_bar.setValue(90)
            
            if resp.status_code >= 400:
                raise Exception(f"HTTP {resp.status_code}: {resp.text}")
            
            self.progress_bar.setValue(100)
            self.log_text.append("✓ 上传完成")
            QMessageBox.information(self, "成功", "指标已成功上传到服务器")
            
        except Exception as e:
            self.log_text.append(f"✗ 上传失败: {e}")
            QMessageBox.critical(self, "错误", str(e))
        finally:
            self.upload_btn.setEnabled(True)


class ChartTab(QWidget):
    """图表生成标签页"""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.init_ui()
    
    def init_ui(self):
        layout = QVBoxLayout()
        
        # 输入区域
        input_group = QGroupBox("数据源")
        input_layout = QVBoxLayout()
        
        file_layout = QHBoxLayout()
        self.data_file_input = QLineEdit()
        browse_btn = QPushButton("浏览...")
        browse_btn.clicked.connect(self.browse_data_file)
        file_layout.addWidget(self.data_file_input)
        file_layout.addWidget(browse_btn)
        input_layout.addLayout(file_layout)
        
        input_group.setLayout(input_layout)
        layout.addWidget(input_group)
        
        # 配置区域
        config_group = QGroupBox("图表配置")
        config_layout = QFormLayout()
        
        output_layout = QHBoxLayout()
        self.output_dir_input = QLineEdit("charts")
        output_browse_btn = QPushButton("浏览...")
        output_browse_btn.clicked.connect(self.browse_output_dir)
        output_layout.addWidget(self.output_dir_input)
        output_layout.addWidget(output_browse_btn)
        config_layout.addRow("输出目录:", output_layout)
        
        self.signal_columns_input = QLineEdit()
        self.signal_columns_input.setPlaceholderText("留空则自动选择前5列，或用逗号分隔")
        config_layout.addRow("信号列:", self.signal_columns_input)
        
        config_group.setLayout(config_layout)
        layout.addWidget(config_group)
        
        # 操作按钮
        self.generate_btn = QPushButton("生成图表")
        self.generate_btn.clicked.connect(self.generate_charts)
        layout.addWidget(self.generate_btn)
        
        # 进度条
        self.progress_bar = QProgressBar()
        layout.addWidget(self.progress_bar)
        
        # 日志输出
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setMaximumHeight(150)
        layout.addWidget(QLabel("日志:"))
        layout.addWidget(self.log_text)
        
        layout.addStretch()
        self.setLayout(layout)
    
    def browse_data_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "选择数据文件", "", 
            "数据文件 (*.parquet *.csv);;所有文件 (*.*)"
        )
        if file_path:
            self.data_file_input.setText(file_path)
    
    def browse_output_dir(self):
        directory = QFileDialog.getExistingDirectory(self, "选择输出目录")
        if directory:
            self.output_dir_input.setText(directory)
    
    def generate_charts(self):
        import pandas as pd
        import matplotlib.pyplot as plt
        
        data_path = Path(self.data_file_input.text())
        output_dir = Path(self.output_dir_input.text())
        signal_columns = self.signal_columns_input.text().strip()
        
        if not data_path.exists():
            QMessageBox.warning(self, "输入错误", "请选择有效的数据文件")
            return
        
        self.generate_btn.setEnabled(False)
        self.log_text.clear()
        self.progress_bar.setValue(0)
        
        try:
            self.log_text.append(f"正在加载数据: {data_path.name}")
            
            if data_path.suffix == ".parquet":
                df = pd.read_parquet(data_path)
            elif data_path.suffix == ".csv":
                df = pd.read_csv(data_path)
            else:
                raise ValueError("不支持的文件类型")
            
            self.progress_bar.setValue(20)
            
            columns = signal_columns.split(",") if signal_columns else df.columns[1:6].tolist()
            output_dir.mkdir(parents=True, exist_ok=True)
            
            total = len(columns)
            for idx, col in enumerate(columns):
                if col not in df.columns:
                    self.log_text.append(f"⚠ 列 {col} 不存在，跳过")
                    continue
                
                plt.figure(figsize=(12, 6))
                plt.plot(df.index if 'timestamps' not in df.columns else df['timestamps'], df[col])
                plt.title(f"Signal: {col}")
                plt.xlabel("Time")
                plt.ylabel(col)
                plt.grid(True, alpha=0.3)
                
                chart_path = output_dir / f"{col}.png"
                plt.savefig(chart_path, dpi=150, bbox_inches='tight')
                plt.close()
                
                self.log_text.append(f"✓ 生成: {chart_path.name}")
                self.progress_bar.setValue(20 + int((idx + 1) * 80 / total))
            
            self.log_text.append(f"\n✓ 所有图表已保存到: {output_dir}")
            QMessageBox.information(self, "成功", f"图表已保存到:\n{output_dir}")
            
        except Exception as e:
            self.log_text.append(f"✗ 错误: {e}")
            QMessageBox.critical(self, "错误", str(e))
        finally:
            self.generate_btn.setEnabled(True)


class ReportTab(QWidget):
    """报表生成标签页"""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.init_ui()
    
    def init_ui(self):
        layout = QVBoxLayout()
        
        # 输入区域
        input_group = QGroupBox("数据源")
        input_layout = QFormLayout()
        
        metrics_layout = QHBoxLayout()
        self.metrics_file_input = QLineEdit()
        metrics_browse_btn = QPushButton("浏览...")
        metrics_browse_btn.clicked.connect(self.browse_metrics_file)
        metrics_layout.addWidget(self.metrics_file_input)
        metrics_layout.addWidget(metrics_browse_btn)
        input_layout.addRow("指标文件:", metrics_layout)
        
        charts_layout = QHBoxLayout()
        self.charts_dir_input = QLineEdit("charts")
        charts_browse_btn = QPushButton("浏览...")
        charts_browse_btn.clicked.connect(self.browse_charts_dir)
        charts_layout.addWidget(self.charts_dir_input)
        charts_layout.addWidget(charts_browse_btn)
        input_layout.addRow("图表目录:", charts_layout)
        
        output_layout = QHBoxLayout()
        self.output_file_input = QLineEdit("report/analysis_report.docx")
        output_browse_btn = QPushButton("浏览...")
        output_browse_btn.clicked.connect(self.browse_output_file)
        output_layout.addWidget(self.output_file_input)
        output_layout.addWidget(output_browse_btn)
        input_layout.addRow("输出文件:", output_layout)
        
        input_group.setLayout(input_layout)
        layout.addWidget(input_group)
        
        # 操作按钮
        self.generate_btn = QPushButton("生成报表")
        self.generate_btn.clicked.connect(self.generate_report)
        layout.addWidget(self.generate_btn)
        
        # 进度条
        self.progress_bar = QProgressBar()
        layout.addWidget(self.progress_bar)
        
        # 日志输出
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setMaximumHeight(150)
        layout.addWidget(QLabel("日志:"))
        layout.addWidget(self.log_text)
        
        layout.addStretch()
        self.setLayout(layout)
    
    def browse_metrics_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "选择指标文件", "", "JSON 文件 (*.json)"
        )
        if file_path:
            self.metrics_file_input.setText(file_path)
    
    def browse_charts_dir(self):
        directory = QFileDialog.getExistingDirectory(self, "选择图表目录")
        if directory:
            self.charts_dir_input.setText(directory)
    
    def browse_output_file(self):
        file_path, _ = QFileDialog.getSaveFileName(
            self, "选择输出文件", "", "Word 文档 (*.docx)"
        )
        if file_path:
            self.output_file_input.setText(file_path)
    
    def generate_report(self):
        from docx import Document
        from docx.shared import Inches
        
        metrics_path = Path(self.metrics_file_input.text())
        charts_dir = Path(self.charts_dir_input.text())
        output_path = Path(self.output_file_input.text())
        
        if not metrics_path.exists():
            QMessageBox.warning(self, "输入错误", "请选择有效的指标文件")
            return
        
        self.generate_btn.setEnabled(False)
        self.log_text.clear()
        self.progress_bar.setValue(0)
        
        try:
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            self.log_text.append("正在创建 Word 文档...")
            doc = Document()
            doc.add_heading("CAN 数据分析报告", 0)
            self.progress_bar.setValue(20)
            
            # 添加指标部分
            doc.add_heading("指标汇总", 1)
            metrics = json.loads(metrics_path.read_text(encoding="utf-8"))
            for key, value in metrics.items():
                doc.add_paragraph(f"{key}: {value}")
            self.progress_bar.setValue(50)
            
            # 添加图表部分
            doc.add_heading("信号可视化", 1)
            if charts_dir.exists():
                chart_files = sorted(charts_dir.glob("*.png"))
                total = len(chart_files)
                for idx, chart in enumerate(chart_files):
                    doc.add_heading(chart.stem, 2)
                    doc.add_picture(str(chart), width=Inches(6))
                    self.log_text.append(f"✓ 添加图表: {chart.name}")
                    self.progress_bar.setValue(50 + int((idx + 1) * 40 / max(total, 1)))
            
            doc.save(str(output_path))
            self.progress_bar.setValue(100)
            self.log_text.append(f"\n✓ 报表已保存到: {output_path}")
            QMessageBox.information(self, "成功", f"报表已生成:\n{output_path}")
            
        except Exception as e:
            self.log_text.append(f"✗ 错误: {e}")
            QMessageBox.critical(self, "错误", str(e))
        finally:
            self.generate_btn.setEnabled(True)


class MainWindow(QMainWindow):
    """主窗口"""
    def __init__(self):
        super().__init__()
        self.init_ui()
    
    def init_ui(self):
        self.setWindowTitle("CAN 数据分析工具 v1.0")
        self.setGeometry(100, 100, 900, 700)
        
        # 创建中心部件
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # 主布局
        main_layout = QVBoxLayout()
        
        # 标题
        title = QLabel("CAN 数据分析工具")
        title_font = QFont()
        title_font.setPointSize(16)
        title_font.setBold(True)
        title.setFont(title_font)
        title.setAlignment(Qt.AlignCenter)
        main_layout.addWidget(title)
        
        # 标签页
        tab_widget = QTabWidget()
        tab_widget.addTab(DownloadTab(), "📥 下载")
        tab_widget.addTab(ComputeTab(), "⚙️ 计算")
        tab_widget.addTab(UploadTab(), "📤 上传")
        tab_widget.addTab(ChartTab(), "📊 图表")
        tab_widget.addTab(ReportTab(), "📄 报表")
        
        main_layout.addWidget(tab_widget)
        
        # 底部信息
        footer = QLabel("© 2026 CAN Data Analysis Platform")
        footer.setAlignment(Qt.AlignCenter)
        footer.setStyleSheet("color: gray; font-size: 10px;")
        main_layout.addWidget(footer)
        
        central_widget.setLayout(main_layout)
        
        # 状态栏
        self.statusBar().showMessage("就绪")


def main():
    app = QApplication(sys.argv)
    app.setApplicationName("CAN Data Analyzer")
    
    window = MainWindow()
    window.show()
    
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
